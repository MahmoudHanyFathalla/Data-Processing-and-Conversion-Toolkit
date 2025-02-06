import pandas as pd
from docx import Document

# Load the Excel file
df = pd.read_excel('C:\\Users\\hp\\Desktop\\Neqabty\\Codes\\hey\\HelpfullScripts\\input\\qfa.xlsx')

# Create a new Word document
doc = Document()

# Loop through each row in the DataFrame
for index, row in df.iterrows():
    # Add the question to the document
    doc.add_paragraph(row['Q'])
    # Add the answer to the document
    doc.add_paragraph(row['A'])
    # Add a blank line for separation
    doc.add_paragraph("\n")

# Save the document
output_path = 'C:\\Users\\hp\\Desktop\\Neqabty\\Codes\\hey\\HelpfullScripts\\input\\output_with_formatting.docx'
doc.save(output_path)

output_path
